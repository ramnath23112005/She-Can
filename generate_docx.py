from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── Page Setup ───────────────────────────────────────────────────────────
for section in doc.sections:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Style Helpers ──────────────────────────────────────────────────────────
PURPLE = RGBColor(0x8B, 0x5C, 0xF6)
PINK = RGBColor(0xEC, 0x48, 0x99)
DARK = RGBColor(0x1F, 0x1F, 0x2E)
GRAY = RGBColor(0x4A, 0x4A, 0x6A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF8, 0xF7, 0xFF)
GREEN = RGBColor(0x10, 0xB9, 0x81)
ORANGE = RGBColor(0xF5, 0x9E, 0x0B)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = PURPLE
            run.font.size = Pt(24)
        elif level == 2:
            run.font.color.rgb = PINK
            run.font.size = Pt(18)
        elif level == 3:
            run.font.color.rgb = DARK
            run.font.size = Pt(14)
    return h

def add_para(text, bold=False, italic=False, color=DARK, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.color.rgb = DARK
        run.font.size = Pt(11)
        run2 = p.add_run(text)
        run2.font.color.rgb = GRAY
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.color.rgb = GRAY
        run.font.size = Pt(11)
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    return p

def shade_cells(row, color_hex):
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=DARK, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = 'Calibri'

# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("✦ She Can Foundation")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = PURPLE
run.font.name = 'Calibri'

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Empowering youth through education,\nopportunities, and digital initiatives.")
run.font.size = Pt(16)
run.font.color.rgb = PINK
run.italic = True
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Full Stack Web Application\nProject Documentation")
run.font.size = Pt(14)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

doc.add_paragraph()

tech = doc.add_paragraph()
tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tech.add_run("Tech Stack: React · Node.js · Express · MongoDB")
run.font.size = Pt(12)
run.font.color.rgb = DARK
run.bold = True
run.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("Table of Contents", level=1)
toc_items = [
    "1. Project Overview",
    "2. Features",
    "3. Tech Stack",
    "4. Folder Structure",
    "5. Backend Architecture",
    "6. Frontend Components",
    "7. API Endpoints",
    "8. Database Schema",
    "9. Installation Guide",
    "10. Deployment Guide",
    "11. Design System",
    "12. Excel Recording System",
    "13. Conclusion",
]
for item in toc_items:
    add_para(item, color=DARK, size=12, space_after=4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("1. Project Overview", level=1)

add_para(
    "She Can Foundation is a modern, full-stack NGO website built to showcase "
    "the mission and impact of a youth-driven non-profit organization focused on "
    "education, women empowerment, community support, and digital awareness. "
    "The project demonstrates professional-grade web development practices and "
    "serves as a compelling portfolio piece for tech internship applications.",
    size=11, color=DARK
)

add_heading_styled("Purpose", level=3)
add_bullet("Present the NGO's mission and core initiatives in an engaging digital experience")
add_bullet("Provide a contact/volunteer form with backend data persistence and email notifications")
add_bullet("Demonstrate full-stack development skills across React, Node.js, Express, and MongoDB")
add_bullet("Showcase modern UI/UX with glassmorphism, dark mode, animations, and responsive design")

add_heading_styled("Target Audience", level=3)
add_bullet("NGO visitors and potential volunteers")
add_bullet("Donors and community partners")
add_bullet("Tech internship evaluators assessing full-stack competency")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. FEATURES
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("2. Features", level=1)

features = [
    ("Landing Page",
     "Modern hero section with animated gradient shapes, floating background blobs, "
     "an inspiring tagline, a badge highlighting 'Youth-Driven NGO', and live stat counters "
     "(500+ Youth Empowered, 25+ Programs, 10+ Communities). Smooth scrolling navbar "
     "with glassmorphism effect on scroll and fully responsive layout."),
    ("About Section",
     "Explains the NGO's mission with a compelling 'Our Story' narrative inside a glass "
     "card. Four feature cards arranged in a 2×2 grid covering Education, Women Empowerment, "
     "Community Support, and Digital Awareness — each with a colored icon and detailed description."),
    ("Our Causes Section",
     "Four gradient-coded cause cards (Quality Education, Digital Literacy, Sustainable Development, "
     "Youth Leadership) with colored headers, impact statistics, and 'Get Involved' CTA links. "
     "Hover animations with card lift and shadow effects."),
    ("Contact / Volunteer Form",
     "Fully validated form with Full Name, Email, Phone Number, and Message fields. "
     "Client-side validation with visual error indicators. Server-side validation using "
     "express-validator. Animated submit button with loading spinner. Success popup with "
     "checkmark animation after submission. Email notification sent to admin via Nodemailer. "
     "Every submission is automatically recorded to submissions_record.xlsx with columns: "
     "Full Name, Email, Phone, Message, Submitted At."),
    ("Dark / Light Mode",
     "Persistent theme toggle stored in localStorage. Smooth CSS transitions across all elements. "
     "Separate CSS variable sets for both themes — light (#ffffff bg) and dark (#0f0f1a bg). "
     "Toggle button rotates on click for a polished feel."),
    ("Scroll Animations",
     "AOS (Animate on Scroll) library integrated. Sections fade and slide up as users scroll. "
     "Staggered animation delays for multi-element layouts (cards, stats). Duration and easing "
     "tuned for a smooth, professional experience."),
    ("Glassmorphism UI",
     "Frosted glass effect using backdrop-filter: blur(16px) with semi-transparent backgrounds "
     "and subtle borders. Applied to cards, navbar, contact form, and submission cards. "
     "Creates a modern, layered visual hierarchy."),
    ("Admin Dashboard",
     "Located at /admin route. Fetches and displays all contact form submissions from the backend. "
     "Each submission card shows user avatar (initials), name, email, phone, timestamp, and message. "
     "Refresh button to reload data. Clean, responsive card layout."),
    ("Excel Data Recording",
     "All contact form submissions are automatically recorded to submissions_record.xlsx "
     "with Full Name, Email, Phone, Message, and timestamp. Newsletter subscriptions are "
     "recorded to a separate newsletter_subscribers.xlsx with Email and timestamp. "
     "Uses the exceljs library. Files are created on first submission and appended to thereafter."),
    ("Loading Animation",
     "Full-screen intro loading animation on page load. Pulsing logo icon with a gradient "
     "progress bar. Auto-dismisses after 2 seconds. Branded with the foundation's visual identity."),
    ("Footer & Newsletter Subscription",
     "Wave SVG divider at the top. Three-column grid: brand/mission with social media icons "
     "(Facebook, Twitter, Instagram, LinkedIn, YouTube), quick links, and newsletter subscription form. "
     "The newsletter form is fully functional — on submit, it saves the email to MongoDB, "
     "records it to newsletter_subscribers.xlsx, sends a welcome email to the subscriber, "
     "and notifies the admin. Shows success/error feedback inline."),
]

for title, desc in features:
    add_heading_styled(title, level=2)
    add_para(desc, color=DARK, size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. TECH STACK
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("3. Tech Stack", level=1)

table = doc.add_table(rows=10, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Layer", "Technology", "Purpose"]
data = [
    ["Frontend", "React 18", "UI component library"],
    ["Build Tool", "Vite 5", "Fast dev server & bundler"],
    ["Routing", "React Router v6", "Client-side routing"],
    ["Animations", "AOS 2.3", "Scroll-triggered animations"],
    ["Icons", "Font Awesome 6", "Icon library"],
    ["Backend", "Node.js + Express 4", "REST API server"],
    ["Database", "MongoDB Atlas", "Cloud document database"],
    ["Email", "Nodemailer", "Email notifications"],
    ["Excel", "ExcelJS", "Record data to Excel files"],
]

for i, h in enumerate(headers):
    set_cell_text(table.rows[0].cells[i], h, bold=True, color=WHITE, size=10)
shade_cells(table.rows[0], "8B5CF6")

for row_idx, row_data in enumerate(data, start=1):
    for col_idx, cell_data in enumerate(row_data):
        set_cell_text(table.rows[row_idx].cells[col_idx], cell_data, size=10)
    if row_idx % 2 == 0:
        shade_cells(table.rows[row_idx], "F3E8FF")

doc.add_paragraph()
add_heading_styled("Additional Packages", level=3)
add_bullet("cors — Cross-Origin Resource Sharing", "")
add_bullet("dotenv — Environment variable management", "")
add_bullet("express-validator — Server-side input validation", "")
add_bullet("mongoose — MongoDB ODM", "")
add_bullet("exceljs — Create and append to Excel files", "")
add_bullet("@fortawesome/react-fontawesome — React icon components", "")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. FOLDER STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("4. Folder Structure", level=1)

structure = """she-can-foundation/
├── backend/
│   ├── models/
│   │   ├── Contact.js              # Contact form schema
│   │   └── Newsletter.js           # Newsletter schema
│   ├── routes/
│   │   ├── contact.js              # Contact API routes
│   │   └── newsletter.js           # Newsletter API routes
│   ├── utils/
│   │   └── excelRecorder.js        # Excel file recording utility
│   ├── server.js                   # Express entry point
│   ├── .env                        # Environment variables
│   └── package.json
├── frontend/
│   ├── public/
│   │   └── favicon.svg             # Brand favicon
│   ├── src/
│   │   ├── components/
│   │   │   ├── Navbar.jsx          # Sticky navbar
│   │   │   ├── Hero.jsx            # Landing hero
│   │   │   ├── About.jsx           # Mission + cards
│   │   │   ├── Causes.jsx          # Cause cards
│   │   │   ├── Contact.jsx         # Contact form
│   │   │   ├── Footer.jsx          # Site footer
│   │   │   ├── Admin.jsx           # Admin dashboard
│   │   │   └── Loading.jsx         # Loading screen
│   │   ├── App.jsx                 # Root component
│   │   ├── App.css                 # All styles
│   │   ├── index.css               # CSS variables
│   │   └── main.jsx                # Entry point
│   ├── index.html
│   ├── vite.config.js
│   ├── vercel.json
│   └── package.json
└── README.md"""

add_code_block(structure)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. BACKEND ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("5. Backend Architecture", level=1)

add_heading_styled("server.js — Entry Point", level=2)
add_para(
    "The Express server initializes CORS, JSON body parsing, and connects to MongoDB "
    "via Mongoose. Routes are mounted under /api. The server listens on the port specified "
    "in the .env file (default 5000). MongoDB connection string is loaded from environment variables.",
    size=11
)
add_code_block(
    "app.use('/api', contactRoutes);\n"
    "mongoose.connect(process.env.MONGODB_URI)\n"
    "  .then(() => app.listen(PORT))\n"
    "  .catch(err => console.error(err))"
)

add_heading_styled("routes/contact.js — API Routes", level=2)
add_para(
    "Two endpoints are defined:\n\n"
    "POST /api/contact — Validates input with express-validator, saves to MongoDB, "
    "and sends an email notification to the admin using Nodemailer with Gmail SMTP.\n\n"
    "GET /api/submissions — Returns all contact form submissions sorted by most recent, "
    "used by the admin dashboard.",
    size=11
)

add_heading_styled("routes/newsletter.js — Newsletter Route", level=2)
add_para(
    "POST /api/newsletter — Validates the email, saves to MongoDB, records to "
    "newsletter_subscribers.xlsx via ExcelJS, sends a welcome email to the subscriber, "
    "and sends a notification email to the admin. Uses the same Nodemailer transporter.",
    size=11
)

add_heading_styled("utils/excelRecorder.js — Excel Recording Utility", level=2)
add_para(
    "Two functions:\n\n"
    "recordToExcel({ fullName, email, phone, message }) — Appends a row to submissions_record.xlsx. "
    "Creates the file with headers on first run, then appends rows on each subsequent submission.\n\n"
    "recordNewsletterToExcel({ email }) — Appends a row to newsletter_subscribers.xlsx with "
    "email and timestamp. Uses a separate file from contact submissions.",
    size=11
)

add_heading_styled("models/Contact.js — Contact Schema", level=2)
add_code_block(
    "{\n"
    "  fullName:  String (required, trimmed),\n"
    "  email:     String (required, lowercase),\n"
    "  phone:     String (required),\n"
    "  message:   String (required),\n"
    "  createdAt: Date (default: Date.now)\n"
    "}"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. FRONTEND COMPONENTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("6. Frontend Components", level=1)

components = [
    ("Navbar (Navbar.jsx)",
     "Fixed-position navigation bar with glassmorphism background on scroll. "
     "Contains the logo (✦ She Can Foundation), four nav links (Home, About, Our Causes, Contact), "
     "an Admin link with outlined styling, and a theme toggle button. On mobile, renders a "
     "slide-in hamburger menu. Active link underline animation on hover."),
    ("Hero (Hero.jsx)",
     "Full-viewport landing section with three floating gradient blob shapes in the background "
     "(purple, pink, light pink) animated with a float keyframe. Displays a badge, the main title "
     "with gradient text effect, a subtitle describing the NGO mission, two CTA buttons "
     "(Explore Our Mission, Join Us Today), and three live stat counters."),
    ("About (About.jsx)",
     "Features a 'Our Story' glass card with the NGO's founding narrative. Below it, a 2×2 grid "
     "of glass cards for Education, Women Empowerment, Community Support, and Digital Awareness. "
     "Each card has a colored icon with background tint, title, and detailed description. "
     "Cards lift on hover with shadow enhancement."),
    ("Causes (Causes.jsx)",
     "Four gradient-coded cause cards in a 2×2 grid. Each card has a colored header with icon "
     "and impact statistic, followed by a body with title, description, and 'Get Involved' link. "
     "Gradients used: Purple, Blue, Green, Orange. Cards lift on hover."),
    ("Contact (Contact.jsx)",
     "Two-column layout: left side has contact info (address, email, phone) with gradient icon "
     "containers; right side is a glassmorphism form with Full Name, Email, Phone, and Message "
     "fields. Real-time client-side validation with regex patterns. Submit button shows a spinner "
     "during submission. On success, a green animated popup appears with checkmark icon. "
     "Fetches POST /api/contact on submit."),
    ("Footer (Footer.jsx)",
     "Wave SVG divider at the top using a path curve. Three-column grid: brand section with "
     "mission statement and 5 social media icon links (Facebook, Twitter, Instagram, LinkedIn, "
     "YouTube); Quick Links section; Newsletter section with email input and subscribe button. "
     "Copyright with current year and heart icon."),
    ("Admin (Admin.jsx)",
     "Dashboard page at /admin route. Fetches GET /api/submissions on mount. Displays a list "
     "of submission cards, each showing a gradient avatar with initials, name, email, phone, "
     "formatted date, and message. Includes a Refresh button with spinning icon. Handles loading, "
     "error, and empty states gracefully."),
    ("Loading (Loading.jsx)",
     "Full-screen centered overlay with a pulsing logo (✦), gradient progress bar animating "
     "left-to-right, and 'She Can Foundation' text. Auto-dismisses after 2 seconds via the "
     "App component."),
]

for title, desc in components:
    add_heading_styled(title, level=2)
    add_para(desc, size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. API ENDPOINTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("7. API Endpoints", level=1)

table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

api_headers = ["Method", "Endpoint", "Description", "Auth"]
api_data = [
    ["POST", "/api/contact", "Submit a contact/volunteer form", "No"],
    ["GET",  "/api/submissions", "Fetch all submissions (admin)", "No"],
    ["POST", "/api/newsletter", "Subscribe to newsletter", "No"],
]

for i, h in enumerate(api_headers):
    set_cell_text(table2.rows[0].cells[i], h, bold=True, color=WHITE, size=10)
shade_cells(table2.rows[0], "8B5CF6")

for row_idx, row_data in enumerate(api_data, start=1):
    for col_idx, cell_data in enumerate(row_data):
        set_cell_text(table2.rows[row_idx].cells[col_idx], cell_data, size=10)
    shade_cells(table2.rows[row_idx], "F3E8FF")

doc.add_paragraph()
add_heading_styled("POST /api/contact — Request Body", level=2)
add_code_block(
    '{\n'
    '  "fullName": "Jane Doe",\n'
    '  "email": "jane@example.com",\n'
    '  "phone": "+91 98765 43210",\n'
    '  "message": "I want to volunteer!"\n'
    '}'
)

add_heading_styled("Response", level=2)
add_code_block(
    '201 Created\n'
    '{\n'
    '  "message": "Thank you! We will get back to you soon."\n'
    '}\n'
    '\n'
    '400 Bad Request (validation error)\n'
    '{\n'
    '  "errors": [\n'
    '    { "msg": "Invalid email address", "path": "email" }\n'
    '  ]\n'
    '}'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("POST /api/newsletter — Request Body", level=2)
add_code_block(
    '{\n'
    '  "email": "jane@example.com"\n'
    '}'
)

add_heading_styled("Response", level=2)
add_code_block(
    '201 Created\n'
    '{\n'
    '  "message": "Subscribed successfully!"\n'
    '}\n'
    '\n'
    '400 Bad Request (validation error)\n'
    '{\n'
    '  "errors": [\n'
    '    { "msg": "Invalid email address", "path": "email" }\n'
    '  ]\n'
    '}'
)

doc.add_page_break()

add_heading_styled("8. Database Schema", level=1)

add_para(
    "The application uses two MongoDB collections within the 'she-can-foundation' database. "
    "Both schemas are defined using Mongoose.",
    size=11
)

add_heading_styled("Collection 1: contacts", level=2)

table3 = doc.add_table(rows=6, cols=4)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

schema_headers = ["Field", "Type", "Required", "Notes"]
schema_data = [
    ["fullName", "String", "Yes", "Trimmed, minLength not enforced"],
    ["email", "String", "Yes", "Trimmed, lowercase, valid email"],
    ["phone", "String", "Yes", "Trimmed, 7-15 digits"],
    ["message", "String", "Yes", "Trimmed, free text"],
    ["createdAt", "Date", "Auto", "Defaults to Date.now"],
]

for i, h in enumerate(schema_headers):
    set_cell_text(table3.rows[0].cells[i], h, bold=True, color=WHITE, size=10)
shade_cells(table3.rows[0], "8B5CF6")

for row_idx, row_data in enumerate(schema_data, start=1):
    for col_idx, cell_data in enumerate(row_data):
        set_cell_text(table3.rows[row_idx].cells[col_idx], cell_data, size=10)
    if row_idx % 2 == 0:
        shade_cells(table3.rows[row_idx], "F3E8FF")

doc.add_paragraph()
add_heading_styled("Mongoose Model Definition", level=2)
add_code_block(
    "const contactSchema = new mongoose.Schema({\n"
    "  fullName: { type: String, required: true, trim: true },\n"
    "  email:    { type: String, required: true, trim: true, lowercase: true },\n"
    "  phone:    { type: String, required: true, trim: true },\n"
    "  message:  { type: String, required: true, trim: true },\n"
    "  createdAt: { type: Date, default: Date.now },\n"
    "});"
)

doc.add_paragraph()

add_heading_styled("Collection 2: newsletters", level=2)

table_nl = doc.add_table(rows=3, cols=4)
table_nl.style = 'Table Grid'
table_nl.alignment = WD_TABLE_ALIGNMENT.CENTER

schema_nl_headers = ["Field", "Type", "Required", "Notes"]
schema_nl_data = [
    ["email", "String", "Yes", "Trimmed, lowercase, valid email"],
    ["subscribedAt", "Date", "Auto", "Defaults to Date.now"],
]

for i, h in enumerate(schema_nl_headers):
    set_cell_text(table_nl.rows[0].cells[i], h, bold=True, color=WHITE, size=10)
shade_cells(table_nl.rows[0], "8B5CF6")

for row_idx, row_data in enumerate(schema_nl_data, start=1):
    for col_idx, cell_data in enumerate(row_data):
        set_cell_text(table_nl.rows[row_idx].cells[col_idx], cell_data, size=10)
    if row_idx % 2 == 0:
        shade_cells(table_nl.rows[row_idx], "F3E8FF")

doc.add_paragraph()
add_heading_styled("Mongoose Model Definition", level=2)
add_code_block(
    "const newsletterSchema = new mongoose.Schema({\n"
    "  email: { type: String, required: true, trim: true, lowercase: true },\n"
    "  subscribedAt: { type: Date, default: Date.now },\n"
    "});"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 9. INSTALLATION GUIDE
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("9. Installation Guide", level=1)

add_heading_styled("Prerequisites", level=2)
add_bullet("Node.js v18+")
add_bullet("npm or yarn")
add_bullet("MongoDB Atlas account (free tier)")
add_bullet("Git")

add_heading_styled("Step 1: Clone the Repository", level=2)
add_code_block("git clone https://github.com/your-username/she-can-foundation.git\ncd she-can-foundation")

add_heading_styled("Step 2: Backend Setup", level=2)
add_code_block(
    "cd backend\n"
    "npm install"
)
add_para("Create a .env file in backend/:")
add_code_block(
    "PORT=5000\n"
    "MONGODB_URI=mongodb+srv://<username>:<password>@cluster0.xxxxx.mongodb.net/she-can-foundation?retryWrites=true&w=majority\n"
    "EMAIL_USER=your-email@gmail.com\n"
    "EMAIL_PASS=your-gmail-app-password\n"
    "ADMIN_EMAIL=admin@shecanfoundation.org"
)
add_para("Start the backend:")
add_code_block("npm run dev")

add_heading_styled("Step 3: Frontend Setup", level=2)
add_code_block(
    "cd frontend\n"
    "npm install\n"
    "npm run dev"
)

add_heading_styled("Step 4: Connect Frontend to Backend", level=2)
add_para(
    "In Contact.jsx and Admin.jsx, update the fetch URLs to point to your backend "
    "(use http://localhost:5000 for local development).",
    size=11
)

add_heading_styled("Quick Start — How to Run", level=2)
add_para("Run the project locally with two terminals:", size=11, bold=True)

add_para("Terminal 1 — Backend:", bold=True, size=11)
add_code_block(
    "cd she-can-foundation/backend\n"
    "npm install\n"
    "# Edit .env with your MongoDB URI\n"
    "npm run dev"
)

add_para("Terminal 2 — Frontend:", bold=True, size=11)
add_code_block(
    "cd she-can-foundation/frontend\n"
    "npm install\n"
    "npm run dev"
)

add_para(
    "The frontend opens at http://localhost:5173 and the backend runs at http://localhost:5000. "
    "Update the fetch URLs in Contact.jsx and Admin.jsx to point to localhost for development.",
    size=11
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 10. DEPLOYMENT GUIDE
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("10. Deployment Guide", level=1)

add_heading_styled("Frontend — Vercel", level=2)
steps_vercel = [
    "Push the project to a GitHub repository.",
    "Go to vercel.com and click 'Add New → Project'.",
    "Import your GitHub repository.",
    "Set the Root Directory to 'frontend'.",
    "Framework Preset will auto-detect 'Vite'.",
    "Click 'Deploy'. The site will be live at a vercel.app URL.",
    "Configure a custom domain in the Vercel dashboard (optional).",
]
for s in steps_vercel:
    add_bullet(s)

add_heading_styled("Backend — Render", level=2)
steps_render = [
    "Go to render.com and click 'New + → Web Service'.",
    "Connect your GitHub repository.",
    "Set the Root Directory to 'backend'.",
    "Start Command: node server.js",
    "Add environment variables from your .env file in the dashboard.",
    "Select the Free tier plan.",
    "Click 'Create Web Service'. The API will be live at a onrender.com URL.",
]
for s in steps_render:
    add_bullet(s)

add_heading_styled("Database — MongoDB Atlas", level=2)
steps_mongo = [
    "Go to mongodb.com/atlas and create a free cluster.",
    "Click 'Connect → Connect your application'.",
    "Copy the connection string.",
    "Replace <username> and <password> with your database user credentials.",
    "In Network Access, add IP 0.0.0.0/0 to allow connections from Render.",
    "Use this string as MONGODB_URI in backend environment variables.",
]
for s in steps_mongo:
    add_bullet(s)

add_heading_styled("Update API URLs for Production", level=2)
add_para(
    "In Contact.jsx and Admin.jsx, replace localhost fetch URLs with your Render backend URL:",
    size=11
)
add_code_block(
    "fetch('https://she-can-foundation-api.onrender.com/api/contact', ...)\n"
    "fetch('https://she-can-foundation-api.onrender.com/api/submissions')"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 11. DESIGN SYSTEM
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("11. Design System", level=1)

add_heading_styled("Color Palette", level=2)

table4 = doc.add_table(rows=5, cols=3)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

color_headers = ["Token", "Hex Value", "Usage"]
color_data = [
    ["Primary", "#8B5CF6", "Buttons, links, gradients, headings"],
    ["Secondary", "#EC4899", "Accents, highlights, badge icons"],
    ["Dark BG", "#0F0F1A", "Dark mode background"],
    ["Light BG", "#FFFFFF", "Light mode background"],
]

for i, h in enumerate(color_headers):
    set_cell_text(table4.rows[0].cells[i], h, bold=True, color=WHITE, size=10)
shade_cells(table4.rows[0], "8B5CF6")

for row_idx, row_data in enumerate(color_data, start=1):
    for col_idx, cell_data in enumerate(row_data):
        set_cell_text(table4.rows[row_idx].cells[col_idx], cell_data, size=10)
    if row_idx % 2 == 0:
        shade_cells(table4.rows[row_idx], "F3E8FF")

doc.add_paragraph()
add_heading_styled("Typography", level=2)
table5 = doc.add_table(rows=3, cols=3)
table5.style = 'Table Grid'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

type_headers = ["Font", "Usage", "Weights"]
type_data = [
    ["Poppins", "Headings (h1-h6)", "400, 500, 600, 700, 800"],
    ["Inter", "Body text, inputs, buttons", "300, 400, 500, 600, 700"],
]

for i, h in enumerate(type_headers):
    set_cell_text(table5.rows[0].cells[i], h, bold=True, color=WHITE, size=10)
shade_cells(table5.rows[0], "8B5CF6")

for row_idx, row_data in enumerate(type_data, start=1):
    for col_idx, cell_data in enumerate(row_data):
        set_cell_text(table5.rows[row_idx].cells[col_idx], cell_data, size=10)
    shade_cells(table5.rows[row_idx], "F3E8FF")

doc.add_paragraph()
add_heading_styled("Glassmorphism Values", level=2)
add_bullet("Background: rgba(255, 255, 255, 0.15) (light), rgba(255, 255, 255, 0.05) (dark)")
add_bullet("Backdrop filter: blur(16px)")
add_bullet("Border: 1px solid rgba(255, 255, 255, 0.25)")
add_bullet("Border radius: 16px")
add_bullet("Box shadow: 0 8px 32px rgba(0, 0, 0, 0.1)")

add_heading_styled("Animations", level=2)
add_bullet("Float: Floating blob shapes with 8s duration, infinite loop")
add_bullet("Loading: Gradient progress bar with scaleX transform, 1.5s loop")
add_bullet("Card hover: translateY(-6px) with shadow enhancement, 0.3s ease")
add_bullet("Navbar underline: width transition from 0 to 100%, 0.3s ease")
add_bullet("Theme toggle: rotate(45deg) on hover, 0.3s ease")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 12. EXCEL RECORDING SYSTEM
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("12. Excel Recording System", level=1)

add_para(
    "As an enhancement beyond standard web application functionality, the system "
    "automatically records all form submissions and newsletter subscriptions to Excel "
    "files for easy offline access, reporting, and data portability.",
    size=11
)

add_heading_styled("submissions_record.xlsx", level=2)
table_ex1 = doc.add_table(rows=2, cols=5)
table_ex1.style = 'Table Grid'
table_ex1.alignment = WD_TABLE_ALIGNMENT.CENTER
ex1_headers = ["Full Name", "Email", "Phone", "Message", "Submitted At"]
for i, h in enumerate(ex1_headers):
    set_cell_text(table_ex1.rows[0].cells[i], h, bold=True, color=WHITE, size=9)
shade_cells(table_ex1.rows[0], "8B5CF6")
sample_data = ["Jane Doe", "jane@example.com", "+91 98765 43210", "I want to volunteer!", "5/27/2026, 10:30 AM"]
for i, val in enumerate(sample_data):
    set_cell_text(table_ex1.rows[1].cells[i], val, size=9)

add_para("")
add_heading_styled("newsletter_subscribers.xlsx", level=2)
table_ex2 = doc.add_table(rows=2, cols=2)
table_ex2.style = 'Table Grid'
table_ex2.alignment = WD_TABLE_ALIGNMENT.CENTER
ex2_headers = ["Email", "Subscribed At"]
for i, h in enumerate(ex2_headers):
    set_cell_text(table_ex2.rows[0].cells[i], h, bold=True, color=WHITE, size=9)
shade_cells(table_ex2.rows[0], "8B5CF6")
set_cell_text(table_ex2.rows[1].cells[0], "jane@example.com", size=9)
set_cell_text(table_ex2.rows[1].cells[1], "5/27/2026, 10:30 AM", size=9)

add_para("")
add_heading_styled("How It Works", level=2)
add_para(
    "The excelRecorder.js utility uses the ExcelJS library to:\n\n"
    "1. Check if the Excel file already exists on disk\n"
    "2. If yes, open it and append a new row\n"
    "3. If no, create a new workbook with a header row and then add the data row\n\n"
    "Both files are stored in the backend/ directory and are generated automatically "
    "on the first form submission or newsletter signup. No manual setup is required.",
    size=11
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 13. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════
add_heading_styled("13. Conclusion", level=1)

add_para(
    "The She Can Foundation website is a production-ready full-stack application that "
    "demonstrates proficiency in modern web development technologies and best practices. "
    "It combines an aesthetically pleasing glassmorphism design with robust backend "
    "functionality, including form validation, database persistence, and email notifications.",
    size=11
)

add_para(
    "Key strengths of the project include:",
    size=11, bold=True
)
add_bullet("Clean separation of concerns with organized frontend/backend folder structure")
add_bullet("Professional UI/UX with dark mode, animations, and responsive design")
add_bullet("Full CRUD capability through REST API endpoints")
add_bullet("Excel data recording for offline record-keeping and reporting")
add_bullet("Newsletter subscription system with automated welcome emails")
add_bullet("Secure input handling with both client and server-side validation")
add_bullet("Comprehensive deployment documentation for Vercel, Render, and MongoDB Atlas")
add_bullet("Beginner-friendly code with clear naming conventions")

doc.add_paragraph()
add_para(
    "This project serves as a compelling portfolio piece for tech internship applications, "
    "showcasing the ability to build, style, and deploy a modern full-stack web application "
    "from scratch.",
    size=11, italic=True, color=GRAY
)

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), "She_Can_Foundation_Documentation.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
